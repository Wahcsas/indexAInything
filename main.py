import docx
import fitz  # PyMuPDF
import re
from docx import Document
from pathlib import Path

# Paths to your files (adjust these as necessary)
docx_path = Path(r"PATH_to_file_with_persons")
pdf_path = Path(r"path_to_book.pdf")

exclude_pages = [1, 2, 3, 4, 5, 6, 7,  31, 32, 33, 34, 35, 58, 79, 105, 139, 140, 141, 163, 164, 165, 196, 197, 198, 222, 223, 257, 258, 259, 260,
                 280, 300, 301, 302, 303, 304, 305, 341, 342, 343, 344, 345, 346, 347, 348, 349, 350, 351]  # pages to exclude, e.g., table of contents
footnote_patterns = [
    r'\d+\t\nSee name+',
    r'\d+\t\nvgl\. name+',
    r'\d+\t\nname+'
]


def extract_names_from_docx(docx_path):
    # Load the DOCX file
    doc = docx.Document(docx_path)
    names = []

    # Assuming each name is in a separate paragraph
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Make sure there's text and not just an empty string
            names.append(text)

    return names


def read_pdf_with_pages(pdf_path):
    # Open the PDF file
    doc = fitz.open(pdf_path)
    pdf_text = {}

    # Extract text from each page
    for page in doc:
        text = page.get_text()
        page_number = page.number + 1  # Page numbers are zero-indexed in PyMuPDF
        pdf_text[page_number] = text

    doc.close()
    return pdf_text


# Extract names and read PDF
names_list = extract_names_from_docx(docx_path)
pdf_pages = read_pdf_with_pages(pdf_path)


def find_name_pages(names, pdf_text, exclude_pages: list, footnote_patterns):
    name_pages = {name: [] for name in names}

    for name in names:
        parts = name.split(', ')
        last_name = parts[0]
        first_name = parts[1] if len(parts) > 1 else ''

        # Pattern to match full name, last name, and possessive forms, but not as part of footnotes
        name_pattern = rf'\b{last_name}(?:,?\s+{first_name})?\b|\b{first_name}\s+{last_name}\b'
        name_regex = re.compile(name_pattern, re.IGNORECASE)

        footnotes_reg_with_name = [n.replace('name', last_name) for n in footnote_patterns]
        # Prepare regex for footnotes
        footnote_regexes = [re.compile(pattern, re.IGNORECASE) for pattern in footnotes_reg_with_name]

        for page_number, text in pdf_text.items():
            if page_number in exclude_pages:
                continue  # Skip excluded pages

            # First check the entire text for the name
            if name_regex.search(text):
                # Check specifically for footnotes
                if any(footnote_regex.search(text) for footnote_regex in footnote_regexes):
                    # If found in footnotes, check if also found outside of the footnotes, by trying to find the
                    # shortest text for each individual split via footnote_regex, we find the text WITHOUT the footnotes
                    text_without_footnotes = min(footnote_regex.split(text)[0] for footnote_regex in footnote_regexes)
                    if name_regex.search(text_without_footnotes):
                        name_pages[name].append(page_number)
                else:
                    name_pages[name].append(page_number)
            # If not found in the main text, don't add the page number even if found in the footnote

    return name_pages


# Use the function to find the pages for each name
name_to_pages = find_name_pages(names=names_list, pdf_text=pdf_pages,
                                exclude_pages=exclude_pages, footnote_patterns=footnote_patterns)

# Print some sample outputs to verify
for name, pages in list(name_to_pages.items())[:5]:  # print results for the first 5 names
    print(f"{name}: {sorted(pages)}")


def update_docx_with_pages(docx_path, name_to_pages, output_path):
    # Load the existing DOCX file
    doc = Document(docx_path)

    # Iterate over each paragraph, check if it's a name, and append page numbers if it is
    for para in doc.paragraphs:
        original_text = para.text.strip()
        if original_text in name_to_pages:
            pages = sorted(name_to_pages[original_text])
            # Create a string of page numbers, separated by commas
            pages_str = ', '.join(map(str, pages))
            # Update the paragraph text with page numbers
            para.text = f"{original_text} {pages_str}"

    # Save the updated document to a new file
    doc.save(output_path)


# Define the path for the output DOCX file
output_docx_path = Path(docx_path.parent, docx_path.stem + '_pages_added' + docx_path.suffix)

# Update the DOCX file with the pages
update_docx_with_pages(docx_path, name_to_pages, output_docx_path)

print("The DOCX file has been updated and saved to:", output_docx_path)
